#!/usr/bin/env python3
"""
Generador de Informes Fondos de Reserva - Dirección Distrital 17D08
Usa Plantilla.docx como base (preserva logo, footer, márgenes exactos).
Novedades 100% dinámicas: título libre + columnas configurables.
"""
import sys, json, re, io, zipfile, os
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import pandas as pd
from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT    = "Aptos"
PT_BODY = 9
PT_TBL  = 7
PT_SM   = 6

NOMINA_HEADERS = [
    "N", "Cedula", "Nombre Afiliado", "Tiene\nSolicitud",
    "Fecha de\nSolicitud", "Tiene Cargos\nPendientes?",
    "Tiene\nderecho", "ESTRUCTURA", "ACUMULA Y/O\nMENSUALIZA"
]
NOMINA_CW = [262, 749, 2583, 638, 789, 817, 595, 789, 626]
EST_ORDER = ["1--11","55-2","55-4","56-1","56-3","57-1","58-1","58-2"]

# ═══════════════════════════════════════════════════════════════
# HELPERS XML — CELDAS Y TABLAS
# ═══════════════════════════════════════════════════════════════
def setup_cell(cell, width, bg=None, valign="center"):
    tc = cell._tc
    old = tc.find(qn("w:tcPr"))
    if old is not None: tc.remove(old)
    pr = OxmlElement("w:tcPr")
    e = OxmlElement("w:tcW"); e.set(qn("w:w"), str(width)); e.set(qn("w:type"), "dxa"); pr.append(e)
    b = OxmlElement("w:tcBorders")
    for s in ["top","left","bottom","right"]:
        se = OxmlElement(f"w:{s}")
        se.set(qn("w:val"), "single"); se.set(qn("w:sz"), "4")
        se.set(qn("w:space"), "0"); se.set(qn("w:color"), "auto"); b.append(se)
    pr.append(b)
    if bg:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), bg)
        pr.append(shd)
    va = OxmlElement("w:vAlign"); va.set(qn("w:val"), valign); pr.append(va)
    tc.insert(0, pr)

def write_cell(cell, text, width, bold=False, italic=False, pt=PT_TBL,
               align="center", bg=None, color=None):
    setup_cell(cell, width, bg=bg)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text) if text is not None else "")
    r.font.name = FONT; r.font.size = Pt(pt); r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    f = OxmlElement("w:rFonts")
    f.set(qn("w:ascii"), FONT); f.set(qn("w:hAnsi"), FONT); f.set(qn("w:cs"), FONT)
    rPr.insert(0, f)

def setup_table(table, col_widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    old_pr = tbl.find(qn("w:tblPr"))
    if old_pr is not None: tbl.remove(old_pr)
    pr = OxmlElement("w:tblPr")
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), str(sum(col_widths))); tw.set(qn("w:type"), "dxa"); pr.append(tw)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pr.append(jc)
    cm = OxmlElement("w:tblCellMar")
    for side, val in [("top","40"),("left","70"),("bottom","40"),("right","70")]:
        s = OxmlElement(f"w:{side}"); s.set(qn("w:w"), val); s.set(qn("w:type"), "dxa"); cm.append(s)
    pr.append(cm)
    tbl.insert(0, pr)
    for old_g in tbl.findall(qn("w:tblGrid")): tbl.remove(old_g)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tbl.insert(1, grid)

# ═══════════════════════════════════════════════════════════════
# HELPERS PÁRRAFOS
# ═══════════════════════════════════════════════════════════════
def _add_run(p, text, pt=PT_BODY, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(pt); r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    f = OxmlElement("w:rFonts")
    f.set(qn("w:ascii"), FONT); f.set(qn("w:hAnsi"), FONT); f.set(qn("w:cs"), FONT)
    rPr.insert(0, f)

def body_para(doc, text, bold=False, italic=False, align="justify", before=60, after=60):
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if align == "justify"
                   else WD_ALIGN_PARAGRAPH.LEFT if align == "left"
                   else WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Twips(before)
    p.paragraph_format.space_after  = Twips(after)
    _add_run(p, text, bold=bold, italic=italic)
    return p

def heading_para(doc, text, underline=False, before=100, after=60):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Twips(before)
    p.paragraph_format.space_after  = Twips(after)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(PT_BODY); r.font.bold = True; r.font.underline = underline
    rPr = r._r.get_or_add_rPr()
    f = OxmlElement("w:rFonts"); f.set(qn("w:ascii"), FONT); f.set(qn("w:hAnsi"), FONT); rPr.insert(0, f)

def spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(40)
    p.paragraph_format.space_after  = Twips(40)

# ═══════════════════════════════════════════════════════════════
# TABLA DATOS GENERALES
# ═══════════════════════════════════════════════════════════════
# Namespaces necesarios para parsear fragmentos de tabla OOXML
_OOXML_NS = (
    'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
    'xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:w10="urn:schemas-microsoft-com:office:word" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
    'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
    'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
)

def _parse_xml_fragment(xml_str):
    """Parsea un fragmento XML OOXML añadiendo los namespaces necesarios."""
    from lxml import etree
    wrapped = f'<_root {_OOXML_NS}>{xml_str}</_root>'
    root = etree.fromstring(wrapped.encode('utf-8'))
    return root[0]  # Primer hijo = el fragmento real

def _replace_wt(xml_str, old_text, new_text):
    """Reemplaza el contenido de un w:t de forma segura."""
    import html as _html
    old_e = _html.escape(str(old_text), quote=False)
    new_e = _html.escape(str(new_text), quote=False)
    return xml_str.replace(f">{old_e}</w:t>", f">{new_e}</w:t>")

def tabla_datos_generales(doc, cfg):
    """
    Inyecta la tabla DATOS GENERALES usando el XML del informe 85 como template.
    Sustituye solo los valores dinámicos preservando toda la estructura exacta
    (vMerge vertical, spans horizontales, estilos, fuentes).
    """
    from pathlib import Path as _Path

    tpl_path = _Path(__file__).parent / "datos_generales_tbl.xml"
    with open(tpl_path, encoding="utf-8") as f:
        xml = f.read()

    mes  = cfg.get("mes", "MES").upper()
    anio = cfg.get("anio", "2026")

    # Sustituir valores dinámicos (texto exacto del informe 85 → nuevo valor)
    xml = _replace_wt(xml, "04-03-2026",                             cfg.get("fecha_informe",""))
    xml = _replace_wt(xml, "DDLC-TH-2026-085",                      cfg.get("nro_informe",""))
    xml = _replace_wt(xml, "Ing. Diana Carolina Sánchez Osorio ", cfg.get("responsable_nombre",""))
    xml = _replace_wt(xml, "3819",                                   cfg.get("responsable_ext",""))
    xml = _replace_wt(xml, "carolina.sanchez@educacion.gob.ec",      cfg.get("responsable_email",""))
    xml = _replace_wt(xml, "JEFE DISTRITAL DE TALENTO HUMANO  ",     cfg.get("responsable_cargo",""))
    xml = _replace_wt(xml, "MSc. Manuel Espinoza Avilés ",            cfg.get("director_nombre",""))
    xml = _replace_wt(xml, "3802",                                   cfg.get("director_ext",""))
    xml = _replace_wt(xml, "manuel.espinozaa@educacion.gob.ec",      cfg.get("director_email",""))
    xml = _replace_wt(xml, "DIRECTOR DISTRITAL 17D08 PARROQUIAS RURALES CONOCOTO LA MERCED",
                            cfg.get("director_cargo",""))
    xml = _replace_wt(xml, "MES DE FEBRERO 2026",                    f"MES DE {mes} {anio}")

    tbl_elem = _parse_xml_fragment(xml)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    body.insert(list(body).index(sectPr), tbl_elem) if sectPr is not None else body.append(tbl_elem)



def tabla_nomina(doc, df_sub, acumula_val, n_start=1):
    """Versión optimizada: construye el XML de toda la tabla en un solo string
    en lugar de llamar write_cell (OxmlElement) por cada celda.
    Con 1200+ filas esto pasa de ~12 minutos a ~20 segundos."""
    from lxml import etree as _et
    CW = NOMINA_CW
    W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    SZ = str(int(PT_TBL * 2))

    def esc(s):
        return (str(s) if s is not None else "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;")

    def tc_xml(text, w, bold=False, bg=None, align="center"):
        shd = f'<w:shd w:val="clear" w:color="auto" w:fill="{bg}"/>' if bg else ""
        b   = "<w:b/><w:bCs/>" if bold else ""
        jc  = "center" if align == "center" else "left"
        return (
            f'<w:tc xmlns:w="{W}">'
            f'<w:tcPr><w:tcW w:w="{w}" w:type="dxa"/>'
            f'<w:tcBorders>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>{shd}'
            f'<w:vAlign w:val="center"/></w:tcPr>'
            f'<w:p><w:pPr><w:jc w:val="{jc}"/></w:pPr>'
            f'<w:r><w:rPr>'
            f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
            f'{b}<w:sz w:val="{SZ}"/><w:szCs w:val="{SZ}"/>'
            f'</w:rPr><w:t xml:space="preserve">{esc(text)}</w:t></w:r>'
            f'</w:p></w:tc>'
        )

    def tr_xml(cells_xml):
        return f'<w:tr xmlns:w="{W}">{"".join(cells_xml)}</w:tr>'

    # Header row
    hdr_cells = [tc_xml(h, w, bold=True, bg="D9D9D9") for h,w in zip(NOMINA_HEADERS, CW)]
    rows_xml = [tr_xml(hdr_cells)]

    # Data rows
    for idx, (_, row) in enumerate(df_sub.iterrows()):
        fecha  = row["FechaSolicitud"] if row["TieneSolicitud"] == "SI" else ""
        ma_val = row.get("AcumulaFinal", acumula_val)
        vals   = [str(n_start+idx), row["Cedula"], row["Nombre"],
                  row["TieneSolicitud"], fecha, row["TieneCargos"],
                  row["TieneDerecho"], row["Estructura"], str(int(ma_val))]
        aligns = ["center","center","left","center","center","center","center","center","center"]
        data_cells = [tc_xml(v, w, align=a) for v,w,a in zip(vals, CW, aligns)]
        rows_xml.append(tr_xml(data_cells))

    # Full table XML
    tw = sum(CW)
    grid = "".join(f'<w:gridCol w:w="{w}"/>' for w in CW)
    tbl_xml = (
        f'<w:tbl xmlns:w="{W}">'
        f'<w:tblPr>'
        f'<w:tblW w:w="{tw}" w:type="dxa"/>'
        f'<w:jc w:val="center"/>'
        f'<w:tblCellMar>'
        f'<w:top w:w="40" w:type="dxa"/><w:left w:w="70" w:type="dxa"/>'
        f'<w:bottom w:w="40" w:type="dxa"/><w:right w:w="70" w:type="dxa"/>'
        f'</w:tblCellMar>'
        f'</w:tblPr>'
        f'<w:tblGrid>{grid}</w:tblGrid>'
        f'{"".join(rows_xml)}'
        f'</w:tbl>'
    )
    tbl_el = _et.fromstring(tbl_xml.encode("utf-8"))
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        body.insert(list(body).index(sectPr), tbl_el)
    else:
        body.append(tbl_el)

# ═══════════════════════════════════════════════════════════════
# TABLA NOVEDAD GENÉRICA (columnas y filas dinámicas)
# ═══════════════════════════════════════════════════════════════
def tabla_novedad(doc, headers, widths, rows):
    """
    headers: lista de strings con nombres de columnas
    widths:  lista de anchos DXA (debe sumar ~8647)
    rows:    lista de listas con los datos
    Si la primera columna es un número secuencial (No, Nro, #), se auto-numera.
    """
    if not rows:
        rows = [["—"] * len(headers)]

    # Detectar si la primera columna es de numeración
    first_h = headers[0].strip().upper().rstrip(".") if headers else ""
    auto_num = first_h in ("NO", "NRO", "#", "N", "NRO.")

    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    setup_table(t, widths)
    hr = t.rows[0]
    for i,(h,w) in enumerate(zip(headers,widths)):
        write_cell(hr.cells[i], h, w, bold=True, bg="D9D9D9")
    for ri, row_data in enumerate(rows):
        dr = t.rows[1+ri]
        for i,(v,w) in enumerate(zip(row_data,widths)):
            # Si es columna de numeración, sobreescribir con número secuencial
            cell_val = str(ri+1) if (auto_num and i == 0) else v
            write_cell(dr.cells[i], cell_val, w, align="center" if i==0 else "left")

# ═══════════════════════════════════════════════════════════════
# TABLA RESUMEN
# ═══════════════════════════════════════════════════════════════
def tabla_resumen(doc, grp_A, grp_B, grp_C, grp_D):
    """
    Inyecta tabla resumen usando el XML de la plantilla oficial (Tablas_Plantilla.docx).
    Sustituye valores por posición (índice de w:t) garantizando distribución idéntica.
    Usa la columna 'Regimen' (LOSEP/CT/LOEI) leída del distributivo.
    """
    from pathlib import Path as _Path

    tpl_path = _Path(__file__).parent / "resumen_tbl.xml"
    with open(tpl_path, encoding="utf-8") as f:
        xml = f.read()

    sin_d = pd.concat([grp_D, grp_A]).copy() if len(grp_A)>0 else grp_D.copy()
    ESTR = [["1--11"],["55-2","55-4"],["56-1","56-3"],["57-1"],["58-1","58-2"]]

    def sv(df, ests, reg=None):
        if reg and "Regimen" in df.columns:
            df = df[df["Regimen"]==reg]
        return [sum(len(df[df["Estructura"]==e]) for e in g) for g in ests]

    # Calcular por régimen real del distributivo
    lsd=sv(sin_d,ESTR,"LOSEP"); lm=sv(grp_C,ESTR,"LOSEP"); la=sv(grp_B,ESTR,"LOSEP")
    lsd_t=sum(lsd); lm_t=sum(lm); la_t=sum(la); lt=lsd_t+lm_t+la_t

    ctsd=sv(sin_d,ESTR,"CT"); ctm=sv(grp_C,ESTR,"CT"); cta=sv(grp_B,ESTR,"CT")
    ctsd_t=sum(ctsd); ctm_t=sum(ctm); cta_t=sum(cta); ctt=ctsd_t+ctm_t+cta_t

    esd=sv(sin_d,ESTR,"LOEI"); em=sv(grp_C,ESTR,"LOEI"); ea=sv(grp_B,ESTR,"LOEI")
    esd_t=sum(esd); em_t=sum(em); ea_t=sum(ea); et=esd_t+em_t+ea_t

    tsd=[lsd[i]+ctsd[i]+esd[i] for i in range(5)]
    tm =[lm[i] +ctm[i] +em[i]  for i in range(5)]
    ta =[la[i] +cta[i] +ea[i]  for i in range(5)]
    tsd_t=lsd_t+ctsd_t+esd_t; tm_t=lm_t+ctm_t+em_t; ta_t=la_t+cta_t+ea_t
    gran=tsd_t+tm_t+ta_t

    def n(v):
        """Formatea número para la tabla — vacío se queda como espacio."""
        return str(v) if v else '0'

    import re as _re

    def replace_nth_wt(xml_in, idx, new_val):
        """Reemplaza el contenido del idx-ésimo w:t en el XML."""
        matches = list(_re.finditer(r'(<w:t[^>]*>)[^<]*(</w:t>)', xml_in))
        if idx >= len(matches):
            return xml_in
        m = matches[idx]
        return xml_in[:m.start()] + m.group(1) + str(new_val) + m.group(2) + xml_in[m.end():]

    # Fila LOSEP (índices 34-52)
    # SD: [34]=1, [35]=55, [36]=56, [37]=57, [38]=58, [39]=TOTAL
    # M:  [40]=1, [41]=55, [42]=56, [43]=57, [44]=58, [45]=TOTAL
    # A:  [46]=1, [47]=55, [48]=56, [49]=57, [50]=58, [51]=TOTAL
    # TOT:[52]
    for i,(idx) in enumerate([34,35,36,37,38]): xml=replace_nth_wt(xml,idx,n(lsd[i]))
    xml=replace_nth_wt(xml,39,n(lsd_t))
    for i,(idx) in enumerate([40,41,42,43,44]): xml=replace_nth_wt(xml,idx,n(lm[i]))
    xml=replace_nth_wt(xml,45,n(lm_t))
    for i,(idx) in enumerate([46,47,48,49,50]): xml=replace_nth_wt(xml,idx,n(la[i]))
    xml=replace_nth_wt(xml,51,n(la_t))
    xml=replace_nth_wt(xml,52,n(lt))

    # Fila CT (índices 54-72)
    for i,(idx) in enumerate([54,55,56,57,58]): xml=replace_nth_wt(xml,idx,n(ctsd[i]))
    xml=replace_nth_wt(xml,59,n(ctsd_t))
    for i,(idx) in enumerate([60,61,62,63,64]): xml=replace_nth_wt(xml,idx,n(ctm[i]))
    xml=replace_nth_wt(xml,65,n(ctm_t))
    for i,(idx) in enumerate([66,67,68,69,70]): xml=replace_nth_wt(xml,idx,n(cta[i]))
    xml=replace_nth_wt(xml,71,n(cta_t))
    xml=replace_nth_wt(xml,72,n(ctt))

    # Fila LOEI (índices 74-92)
    for i,(idx) in enumerate([74,75,76,77,78]): xml=replace_nth_wt(xml,idx,n(esd[i]))
    xml=replace_nth_wt(xml,79,n(esd_t))
    for i,(idx) in enumerate([80,81,82,83,84]): xml=replace_nth_wt(xml,idx,n(em[i]))
    xml=replace_nth_wt(xml,85,n(em_t))
    for i,(idx) in enumerate([86,87,88,89,90]): xml=replace_nth_wt(xml,idx,n(ea[i]))
    xml=replace_nth_wt(xml,91,n(ea_t))
    xml=replace_nth_wt(xml,92,n(et))

    # Fila TOTAL (índices 94-112)
    for i,(idx) in enumerate([94,95,96,97,98]): xml=replace_nth_wt(xml,idx,n(tsd[i]))
    xml=replace_nth_wt(xml,99,n(tsd_t))
    for i,(idx) in enumerate([100,101,102,103,104]): xml=replace_nth_wt(xml,idx,n(tm[i]))
    xml=replace_nth_wt(xml,105,n(tm_t))
    for i,(idx) in enumerate([106,107,108,109,110]): xml=replace_nth_wt(xml,idx,n(ta[i]))
    xml=replace_nth_wt(xml,111,n(ta_t))
    xml=replace_nth_wt(xml,112,n(gran))

    # Total funcionarios (índice 129)
    xml=replace_nth_wt(xml,129,n(gran))

    tbl_elem = _parse_xml_fragment(xml)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    body.insert(list(body).index(sectPr), tbl_elem) if sectPr is not None else body.append(tbl_elem)
    return gran


def tabla_firmas(doc, cfg):
    CW = [2882,2882,2883]
    for nombre,cargo in [
        (cfg.get("elabora_nombre",""), cfg.get("elabora_cargo","")),
        (cfg.get("revisa_nombre",""),  cfg.get("revisa_cargo","")),
    ]:
        t = doc.add_table(rows=3, cols=3); setup_table(t, CW)
        write_cell(t.rows[0].cells[0],"DESARROLLO DEL DOCUMENTO",CW[0],bold=True,bg="D9D9D9")
        write_cell(t.rows[0].cells[1],"",CW[1],bg="D9D9D9")
        write_cell(t.rows[0].cells[2],"",CW[2],bg="D9D9D9")
        write_cell(t.rows[1].cells[0],"Nombre",CW[0],bold=True)
        write_cell(t.rows[1].cells[1],"Firma",CW[1],bold=True)
        write_cell(t.rows[1].cells[2],"Fecha",CW[2],bold=True)
        write_cell(t.rows[2].cells[0],f"{nombre}\n{cargo}",CW[0])
        write_cell(t.rows[2].cells[1],"",CW[1])
        write_cell(t.rows[2].cells[2],cfg.get("fecha_informe",""),CW[2])
        spacer(doc)

# ═══════════════════════════════════════════════════════════════
# DATOS — LECTURA Y CRUCE
# ═══════════════════════════════════════════════════════════════
def decode_est(code):
    s = str(code).strip()
    if len(s) < 25: return s
    prog = s[15:17].lstrip("0") or "0"
    sub  = s[24].lstrip("0") or "0"
    return "1--11" if prog == "1" else f"{prog}-{sub}"

def norm_ced(x):
    try: return str(int(float(str(x).replace(",","")))).strip()
    except: return str(x).strip()

def fmt_fecha(v):
    if v is None or str(v).strip() in ("","NaT","nan","None","NaN"): return ""
    try:
        d = pd.to_datetime(v); return f"{d.day}/{d.month}/{d.year}"
    except: return str(v)

def _excel_engine(path):
    return "xlrd" if str(path).lower().endswith(".xls") else "openpyxl"

def _leer_excel(path, **kwargs):
    """Lee .xls o .xlsx con el motor disponible."""
    import subprocess, tempfile
    from pathlib import Path as _P
    ext = _P(str(path)).suffix.lower()
    if ext == ".xlsx":
        return pd.read_excel(path, engine="openpyxl", **kwargs)
    if ext == ".xls":
        try:
            return pd.read_excel(path, engine="xlrd", **kwargs)
        except ImportError:
            pass
        for cmd in ["soffice","libreoffice","/usr/bin/soffice",
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]:
            try:
                r = subprocess.run([cmd,"--version"],capture_output=True,timeout=5)
                if r.returncode == 0:
                    tmp = tempfile.mkdtemp()
                    subprocess.run([cmd,"--headless","--convert-to","xlsx",
                                   str(path),"--outdir",tmp],capture_output=True,timeout=60)
                    conv = os.path.join(tmp, _P(path).stem+".xlsx")
                    if os.path.exists(conv):
                        return pd.read_excel(conv, engine="openpyxl", **kwargs)
            except Exception:
                continue
        raise ValueError(
            "No se puede leer el archivo .xls.\n"
            "SOLUCIÓN: En Excel, Archivo → Guardar como → Excel (.xlsx)\n"
            "O instale xlrd: pip install xlrd"
        )
    raise ValueError(f"Formato no soportado: {ext}")

def load_data(pp_path, pm_path, pd_path):
    """
    Nueva lógica de clasificación con 3 fuentes:
      pp_path = Planilla de Pago IESS  → cédulas aquí = código 2 (Acumulan), DIRECTO
      pm_path = Matriz de Acumulación  → fuente de TieneSolicitud / TieneDerecho
      pd_path = Distributivo Mineduc   → fuente de Nombre, Estructura, Régimen

    Tabla de verdad (para los NO-acumulan):
      Solicitud=SI|NO + Derecho=NO → 0 (Sin Derecho)
      Solicitud=NO    + Derecho=SI → 1 (Mensualiza)
    """
    # ── Planilla de Pago: cédulas que ACUMULAN (código 2) ─────
    try:
        df_plan_raw = _leer_excel(pp_path, header=None)
    except Exception as e:
        raise ValueError(f"No se pudo leer la Planilla de Pago (Archivo 1): {e}")
    # Buscar fila de cabecera (puede estar en las primeras 15 filas)
    hrow_p = 0
    for i, row in df_plan_raw.head(20).iterrows():
        vals_lower = [str(v).strip().lower() for v in row if str(v).strip() and str(v).strip().lower() != 'nan']
        if any("cedula" in v or "identificacion" in v or "identificación" in v for v in vals_lower):
            hrow_p = i; break
    # Renombrar columnas con los valores del header row encontrado
    raw_headers = [str(c).strip() for c in df_plan_raw.iloc[hrow_p]]
    # Deduplicate headers that might be nan
    seen = {}
    clean_headers = []
    for h in raw_headers:
        if h in seen:
            seen[h] += 1
            clean_headers.append(f"{h}_{seen[h]}" if h != 'nan' else f"_col{len(clean_headers)}")
        else:
            seen[h] = 0
            clean_headers.append(h if h != 'nan' else f"_col{len(clean_headers)}")
    df_plan_raw.columns = clean_headers
    df_plan = df_plan_raw.iloc[hrow_p+1:].reset_index(drop=True)
    # Find cedula column and dropna on it (NOT on column 0 which may be nan)
    col_ced_p = next((c for c in df_plan.columns
                      if "cedula" in c.lower() or "identificacion" in c.lower()
                      or "identificación" in c.lower()), None)
    if col_ced_p is None:
        raise ValueError(f"No se encontró columna CEDULA en la Planilla de Pago.\n"
                         f"Columnas detectadas: {list(df_plan.columns)[:10]}")
    df_plan = df_plan.dropna(subset=[col_ced_p])
    # Filter only rows where cedula looks like a number
    df_plan = df_plan[df_plan[col_ced_p].apply(lambda x: str(x).strip().replace('.','').replace(' ','').isdigit())]
    acumulan_ceds = set(df_plan[col_ced_p].apply(norm_ced).tolist())
    print(f"[LOAD] Planilla de pago: {len(acumulan_ceds)} acumulantes (código 2)")

    # ── Distributivo ─────────────────────────────────────────
    df_dist = _leer_excel(pd_path, header=0, dtype={"ESTRUCTURA PROGRAMÁTICA": str})
    df_dist.columns = [str(c).strip() for c in df_dist.columns]
    required = ["ESTRUCTURA PROGRAMÁTICA","ESTADO DEL PUESTO","NOMBRES","NÚMERO IDENTIFICACIÓN"]
    for col in required:
        if col not in df_dist.columns:
            raise ValueError(f"Columna requerida no encontrada en el Distributivo: '{col}'\n"
                             f"Columnas disponibles: {list(df_dist.columns)[:10]}\n"
                             f"NOTA: Verifique que subió el Distributivo (Mineduc) como Archivo 3.")
    col_regimen = next((c for c in df_dist.columns
                        if "GIMEN" in c.upper() and "LABORAL" in c.upper()
                        and "DIGO" not in c.upper()), None)
    col_cod_reg = next((c for c in df_dist.columns
                        if "DIGO" in c.upper() and "GIMEN" in c.upper()), None)
    occ = {}
    for _, row in df_dist.iterrows():
        if str(row["ESTADO DEL PUESTO"]).upper() == "OCUPADO" and pd.notna(row["NÚMERO IDENTIFICACIÓN"]):
            c   = norm_ced(row["NÚMERO IDENTIFICACIÓN"])
            raw = str(row["ESTRUCTURA PROGRAMÁTICA"]) if pd.notna(row["ESTRUCTURA PROGRAMÁTICA"]) else ""
            reg = "LOEI"
            if col_cod_reg and pd.notna(row.get(col_cod_reg)):
                cod = int(float(str(row[col_cod_reg]))) if str(row[col_cod_reg]).replace('.','').isdigit() else 0
                reg = {1:"LOSEP", 2:"CT", 3:"LOEI"}.get(cod, "LOEI")
            elif col_regimen and pd.notna(row.get(col_regimen)):
                rv = str(row[col_regimen]).upper()
                if "LOSEP" in rv or "CIVIL" in rv: reg = "LOSEP"
                elif "TRABAJO" in rv or "CT" in rv: reg = "CT"
            occ[c] = {
                "nombre":     str(row["NOMBRES"]).strip() if pd.notna(row["NOMBRES"]) else "",
                "estructura": decode_est(raw) if raw else "",
                "regimen":    reg
            }

    # ── Matriz IESS ──────────────────────────────────────────
    try:
        xl = pd.ExcelFile(pm_path, engine=_excel_engine(pm_path))
        # Buscar la hoja que tiene la cabecera Cedula en columnas reales.
        # Se saltan hojas de 1 columna (CSV embebidos) que pueden contener
        # "cedula" concatenado con todo el resto.
        sheet = None
        for sh in xl.sheet_names:
            try:
                sample = xl.parse(sh, header=None, nrows=8)
                if sample.shape[1] < 3:   # hoja CSV de una sola columna — saltar
                    continue
                for _, row in sample.iterrows():
                    if any("cedula" in str(v).lower() for v in row
                           if str(v).strip() and str(v).strip().lower() != "nan"):
                        sheet = sh
                        break
                if sheet:
                    break
            except Exception:
                continue
        if sheet is None:
            sheet = xl.sheet_names[0]
        df = xl.parse(sheet, header=None)
    except Exception as e:
        raise ValueError(f"No se pudo leer la Matriz de Acumulación (Archivo 2): {e}")
    hrow = -1
    for i, row in df.iterrows():
        if any("cedula" in str(v).lower() for v in row if str(v).strip() and str(v).strip().lower() != 'nan'):
            hrow = i; break
    if hrow < 0:
        raise ValueError("No se encontró la fila de cabecera (Cedula) en la Matriz de Acumulación (Archivo 2).")
    df.columns = [str(c).strip() for c in df.iloc[hrow]]
    df = df.iloc[hrow+1:].reset_index(drop=True)
    # Drop based on cedula column, not column[0] which may be nan
    col0 = next((c for c in df.columns if "cedula" in c.lower()), df.columns[0])
    df = df.dropna(subset=[col0])

    # Mapear columnas
    col_map = {}
    for col in df.columns:
        cl = col.lower().strip()
        if "cedula" in cl:                               col_map["Cedula"]         = col
        elif "nombre" in cl:                             col_map["Nombre"]          = col
        elif "solicitud" in cl and "tiene" in cl:        col_map["TieneSolicitud"] = col
        elif "fecha" in cl:                              col_map["FechaSolicitud"]  = col
        elif "cargos" in cl:                             col_map["TieneCargos"]     = col
        elif "derecho" in cl or cl.strip() in ("tiene","tiene "): col_map["TieneDerecho"] = col

    df["Cedula"] = df[col_map["Cedula"]].apply(norm_ced)
    df["Nombre"] = df[col_map.get("Nombre", col_map["Cedula"])].apply(lambda x: str(x).strip())
    df["TieneSolicitud"] = df[col_map["TieneSolicitud"]].apply(lambda x: str(x).strip().upper())
    df["FechaSolicitud"] = df[col_map["FechaSolicitud"]].apply(fmt_fecha) if "FechaSolicitud" in col_map else ""
    df["TieneCargos"]    = df[col_map["TieneCargos"]].apply(lambda x: str(x).strip().upper())
    df["TieneDerecho"]   = df[col_map["TieneDerecho"]].apply(lambda x: str(x).strip().upper())

    # ── NUEVA TABLA DE VERDAD ─────────────────────────────────
    # Paso 1: Cédula en Planilla de Pago → MA=2 (Acumula), DIRECTO, sin más análisis
    # Paso 2: Para el resto → analizar TieneSolicitud / TieneDerecho:
    #   Derecho=NO (con o sin solicitud) → MA=0 (Sin Derecho)
    #   Derecho=SI + Solicitud!=SI       → MA=1 (Mensualiza)
    #   Derecho=SI + Solicitud=SI pero no en planilla → MA=1 (edge case, no debería ocurrir)
    def calcular_ma(row):
        c = row["Cedula"]
        if c in acumulan_ceds:
            return 2  # Planilla de pago: es acumulante
        der = row["TieneDerecho"]
        if der == "NO":
            return 0  # Sin Derecho (Solicitud SI o NO, da igual)
        # Derecho=SI y no está en planilla → Mensualiza
        return 1

    df["MA"] = df.apply(calcular_ma, axis=1)
    print(f"[LOAD] Clasificación: 0(SD)={len(df[df.MA==0])}  1(M)={len(df[df.MA==1])}  2(A)={len(df[df.MA==2])}")

    result, sm = [], []
    for _, r in df.iterrows():
        c = r["Cedula"]
        rec = {"Cedula": c, "TieneSolicitud": r["TieneSolicitud"],
               "FechaSolicitud": r["FechaSolicitud"] if "FechaSolicitud" in col_map else "",
               "TieneCargos": r["TieneCargos"], "TieneDerecho": r["TieneDerecho"],
               "MA": r["MA"]}
        if c in occ:
            rec["Nombre"]    = occ[c]["nombre"]
            rec["Estructura"]= occ[c]["estructura"]
            rec["Regimen"]   = occ[c]["regimen"]
            result.append(rec)
        else:
            rec["Nombre"]    = r["Nombre"]
            rec["Estructura"]= "PENDIENTE"
            rec["Regimen"]   = "LOEI"
            sm.append(rec)
    return pd.DataFrame(result + sm), sm
    df_dist.columns = [str(c).strip() for c in df_dist.columns]
    required = ["ESTRUCTURA PROGRAMÁTICA","ESTADO DEL PUESTO","NOMBRES","NÚMERO IDENTIFICACIÓN"]
    for col in required:
        if col not in df_dist.columns:
            raise ValueError(f"Columna requerida no encontrada en el Distributivo: '{col}'\n"
                             f"Columnas disponibles: {list(df_dist.columns)[:10]}\n"
                             f"NOTA: Verifique que subió el Distributivo (Mineduc) como Archivo 2 y la Matriz IESS como Archivo 1.")
    # Detectar columna de régimen laboral (puede variar el nombre)
    col_regimen = next((c for c in df_dist.columns
                        if "GIMEN" in c.upper() and "LABORAL" in c.upper()
                        and "DIGO" not in c.upper()), None)
    col_cod_reg = next((c for c in df_dist.columns
                        if "DIGO" in c.upper() and "GIMEN" in c.upper()), None)
    occ = {}
    for _, row in df_dist.iterrows():
        if str(row["ESTADO DEL PUESTO"]).upper() == "OCUPADO" and pd.notna(row["NÚMERO IDENTIFICACIÓN"]):
            c   = norm_ced(row["NÚMERO IDENTIFICACIÓN"])
            raw = str(row["ESTRUCTURA PROGRAMÁTICA"]) if pd.notna(row["ESTRUCTURA PROGRAMÁTICA"]) else ""
            # Determinar régimen: 1=LOSEP, 2=CT, 3=LOEI
            reg = "LOEI"  # default
            if col_cod_reg and pd.notna(row.get(col_cod_reg)):
                cod = int(float(str(row[col_cod_reg]))) if str(row[col_cod_reg]).replace('.','').isdigit() else 0
                reg = {1:"LOSEP", 2:"CT", 3:"LOEI"}.get(cod, "LOEI")
            elif col_regimen and pd.notna(row.get(col_regimen)):
                rv = str(row[col_regimen]).upper()
                if "LOSEP" in rv or "CIVIL" in rv: reg = "LOSEP"
                elif "TRABAJO" in rv or "CT" in rv: reg = "CT"
            occ[c] = {
                "nombre":     str(row["NOMBRES"]).strip() if pd.notna(row["NOMBRES"]) else "",
                "estructura": decode_est(raw) if raw else "",
                "regimen":    reg
            }

    # ── Matriz IESS ──────────────────────────────────────────
    eng_m = _excel_engine(pm_path)
    try:
        xl = pd.ExcelFile(pm_path, engine=eng_m)
        sheet = next((s for s in xl.sheet_names if "hoja2" in s.lower()), xl.sheet_names[-1])
        df = xl.parse(sheet, header=None)
    except Exception as e:
        raise ValueError(f"No se pudo leer la Matriz IESS: {e}")
    hrow = -1
    for i, row in df.iterrows():
        if any("cedula" in str(v).lower() for v in row if str(v).strip()):
            hrow = i; break
    if hrow < 0:
        raise ValueError("No se encontró la fila de cabecera (Cedula) en la Matriz IESS.\n"
                         "NOTA: Verifique que subió la Matriz IESS como Archivo 1 y el Distributivo como Archivo 2.")
    df.columns = [str(c).strip() for c in df.iloc[hrow]]
    df = df.iloc[hrow+1:].reset_index(drop=True)
    df = df.dropna(subset=[df.columns[0]])

    # Mapear columnas flexiblemente
    col_map = {}
    col_ma = None
    for col in df.columns:
        cl = col.lower().strip()
        if "cedula" in cl:                               col_map["Cedula"]         = col
        elif "nombre" in cl:                             col_map["Nombre"]          = col
        elif "solicitud" in cl and "tiene" in cl:        col_map["TieneSolicitud"] = col
        elif "fecha" in cl:                              col_map["FechaSolicitud"]  = col
        elif "cargos" in cl:                             col_map["TieneCargos"]     = col
        elif "derecho" in cl or cl.strip() in ("tiene","tiene "): col_map["TieneDerecho"] = col
        elif "mensualiza" in cl or "acumula" in cl:      col_ma = col

    df["Cedula"] = df[col_map["Cedula"]].apply(norm_ced)
    df["Nombre"] = df[col_map.get("Nombre", col_map["Cedula"])].apply(lambda x: str(x).strip())
    df["TieneSolicitud"] = df[col_map["TieneSolicitud"]].apply(lambda x: str(x).strip().upper())
    df["FechaSolicitud"] = df[col_map["FechaSolicitud"]].apply(fmt_fecha) if "FechaSolicitud" in col_map else ""
    df["TieneCargos"]    = df[col_map["TieneCargos"]].apply(lambda x: str(x).strip().upper())
    df["TieneDerecho"]   = df[col_map["TieneDerecho"]].apply(lambda x: str(x).strip().upper())

    # Columna MENSUALIZA Y/O ACUMULA: 0=SD, 1=Mensualiza, 2=Acumula
    if col_ma:
        df["MA"] = df[col_ma].apply(lambda x: int(float(x)) if pd.notna(x) and str(x).strip() else 0)
    else:
        # Fallback: deducir de SI/NO como antes
        df["MA"] = df.apply(lambda r: 0 if r["TieneDerecho"]=="NO"
                            else (2 if r["TieneSolicitud"]=="SI" else 1), axis=1)

    result, sm = [], []
    for _, r in df.iterrows():
        c = r["Cedula"]
        rec = {"Cedula": c, "TieneSolicitud": r["TieneSolicitud"],
               "FechaSolicitud": r["FechaSolicitud"] if "FechaSolicitud" in col_map else "",
               "TieneCargos": r["TieneCargos"], "TieneDerecho": r["TieneDerecho"],
               "MA": r["MA"]}
        if c in occ:
            rec["Nombre"]    = occ[c]["nombre"]
            rec["Estructura"]= occ[c]["estructura"]
            rec["Regimen"]   = occ[c]["regimen"]
            result.append(rec)
        else:
            rec["Nombre"]    = r["Nombre"]
            rec["Estructura"]= "PENDIENTE"
            rec["Regimen"]   = "LOEI"
            sm.append(rec)
    return pd.DataFrame(result + sm), sm

def split_by_est(df):
    out = {}
    for e in EST_ORDER:
        s = df[df["Estructura"]==e]
        if len(s): out[e] = s.reset_index(drop=True)
    for e in df["Estructura"].unique():
        if e not in out and len(df[df["Estructura"]==e]):
            out[e] = df[df["Estructura"]==e].reset_index(drop=True)
    return out

# ═══════════════════════════════════════════════════════════════
# GENERADOR DE EXCEL DE RESPALDO
# ═══════════════════════════════════════════════════════════════
def generar_excel_respaldo(cfg, novedades, sA, sB, sC, sD, gran_total, output_path,
                           grp_A=None, grp_B=None, grp_C=None, grp_D=None,
                           renuncias_df=None, renuncia_extras=None):
    """
    Genera un Excel de respaldo con CADA tabla del informe en su propia hoja
    (una hoja por novedad / categoría / nómina automática) + una hoja RESUMEN
    al final. Mismo orden y datos que el Word; formato por hojas como el
    generador de nómina.
    """
    import re as _re
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)   # se crea una hoja por tabla; no se usa la hoja por defecto

    # Estilos
    hdr_fill   = PatternFill("solid", fgColor="1A3A5C")   # azul oscuro
    sub_fill   = PatternFill("solid", fgColor="D9D9D9")   # gris
    sep_fill   = PatternFill("solid", fgColor="366092")   # azul medio
    hdr_font   = Font(name="Calibri", bold=True, color="FFFFFF", size=9)
    sub_font   = Font(name="Calibri", bold=True, size=8)
    body_font  = Font(name="Calibri", size=8)
    title_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    ws = None
    cur_row = 1
    _used = set()

    def _sheet_name(base):
        # Nombre de hoja válido: <=31 chars, sin / \ ? * [ ] :, único
        name = _re.sub(r'[/\\\?\*\[\]:]', '-', str(base)).strip()[:31] or "Hoja"
        cand = name; i = 1
        while cand.lower() in _used:
            suf = f" ({i})"; cand = name[:31 - len(suf)] + suf; i += 1
        _used.add(cand.lower())
        return cand

    def nueva_hoja(base, ncols=10):
        nonlocal ws, cur_row
        ws = wb.create_sheet(title=_sheet_name(base))
        cur_row = 1
        for col in range(1, max(ncols, 1) + 2):
            ws.column_dimensions[get_column_letter(col)].width = 18

    def write_hdr(row, col, text, fill=None, font=None, align=center):
        c = ws.cell(row=row, column=col, value=text)
        c.fill   = fill or sub_fill
        c.font   = font or sub_font
        c.alignment = align
        c.border = border
        return c

    def write_body(row, col, text, align=center):
        c = ws.cell(row=row, column=col, value=text)
        c.font      = body_font
        c.alignment = align
        c.border    = border
        return c

    def write_section_title(row, text, ncols=10):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        c = ws.cell(row=row, column=1, value=text)
        c.fill      = sep_fill
        c.font      = title_font
        c.alignment = center
        ws.row_dimensions[row].height = 18

    item_num = 5

    # ── Novedades manuales: una hoja cada una ─────────────────
    for nov in novedades:
        titulo  = nov.get("titulo", "NOVEDAD").strip()
        headers = nov.get("headers", [])
        filas   = nov.get("filas", [])
        if not headers:
            continue
        nueva_hoja(titulo or f"Novedad {item_num}", ncols=len(headers))
        write_section_title(cur_row, f"{item_num}.- NOVEDAD: {titulo}", ncols=len(headers))
        cur_row += 1
        for ci, h in enumerate(headers):
            write_hdr(cur_row, ci + 1, h)
        cur_row += 1
        for fila in filas:
            for ci, v in enumerate(fila):
                write_body(cur_row, ci + 1, v, align=(center if ci == 0 else left))
            cur_row += 1
        if not filas:
            ws.merge_cells(start_row=cur_row, start_column=1, end_row=cur_row, end_column=len(headers))
            write_body(cur_row, 1, "(sin registros)")
            cur_row += 1
        item_num += 1

    # ── Desvinculaciones por categoría: una hoja cada una ─────
    if renuncias_df is not None and len(renuncias_df) > 0:
        if renuncia_extras is None:
            renuncia_extras = {}
        ren_hdrs = ["No", "CÉDULA", "APELLIDOS Y NOMBRES", "RMU", "OBSERVACIÓN"]
        CATEGORIAS_SM = ["RENUNCIA", "NOTIFICACIÓN DE SALIDA", "FALLECIMIENTO"]
        for cat in CATEGORIAS_SM:
            cat_rows = []
            for _, row in renuncias_df.iterrows():
                ced = row["Cedula"]
                extras = renuncia_extras.get(ced, {})
                row_cat = extras.get("categoria", "RENUNCIA")
                if row_cat == cat:
                    cat_rows.append((ced, row, extras))
            if not cat_rows:
                continue
            nueva_hoja(cat, ncols=len(ren_hdrs))
            write_section_title(cur_row, f"{item_num}.- NOVEDAD: {cat}", ncols=len(ren_hdrs))
            cur_row += 1
            for ci, h in enumerate(ren_hdrs):
                write_hdr(cur_row, ci + 1, h)
            cur_row += 1
            for i, (ced, row, extras) in enumerate(cat_rows):
                vals = [
                    str(i + 1), ced,
                    row["Nombre"] if row["Nombre"] and str(row["Nombre"]) != "nan" else "",
                    extras.get("rmu", ""),
                    extras.get("observacion", f"{cat} / NO CONSTA EN DISTRIBUTIVO")
                ]
                for ci, v in enumerate(vals):
                    write_body(cur_row, ci + 1, v, align=(center if ci == 0 else left))
                cur_row += 1
            item_num += 1

    # ── Nóminas automáticas: una hoja cada una ────────────────
    NOMINA_HDR = ["N", "Cedula", "Nombre Afiliado", "Tiene Solicitud", "Fecha Solicitud",
                  "Tiene Cargos", "Tiene Derecho", "Estructura", "Acumula"]

    def write_nomina(titulo, grupos, acumula_val, item_n):
        nonlocal cur_row
        nueva_hoja(titulo, ncols=len(NOMINA_HDR))
        write_section_title(cur_row, f"{item_n}. {titulo}", ncols=len(NOMINA_HDR))
        cur_row += 1
        for ci, h in enumerate(NOMINA_HDR):
            write_hdr(cur_row, ci + 1, h)
        cur_row += 1
        counter = 1
        for est, sub in grupos.items():
            for _, row in sub.iterrows():
                fecha = row["FechaSolicitud"] if row["TieneSolicitud"] == "SI" else ""
                vals = [str(counter), row["Cedula"], row["Nombre"],
                        row["TieneSolicitud"], fecha, row["TieneCargos"],
                        row["TieneDerecho"], row["Estructura"], str(acumula_val)]
                for ci, v in enumerate(vals):
                    write_body(cur_row, ci + 1, v, align=(center if ci in [0, 3, 4, 5, 6, 7, 8] else left))
                cur_row += 1
                counter += 1
        if counter == 1:
            write_body(cur_row, 1, "(sin registros)"); cur_row += 1

    write_nomina("14.1 SOLICITUD SI / DERECHO NO", sA, 0, item_num)
    write_nomina("14.1 SOLICITUD NO / DERECHO NO", sD, 0, item_num)
    write_nomina("14.2 ACUMULAN (SI/SI)",           sB, 2, item_num)
    write_nomina("14.3 MENSUALIZA (NO/SI)",         sC, 1, item_num)

    # ── Hoja RESUMEN (al final) ───────────────────────────────
    NCOLS_RES = 8
    nueva_hoja("RESUMEN", ncols=NCOLS_RES)
    write_section_title(cur_row, "TABLA RESUMEN — CANCELACIÓN FONDOS DE RESERVA", ncols=NCOLS_RES)
    cur_row += 1

    if grp_A is not None and grp_B is not None and grp_C is not None and grp_D is not None:
        sin_d = pd.concat([grp_D, grp_A]).copy() if len(grp_A) > 0 else grp_D.copy()
        ESTR = [["1--11"], ["55-2", "55-4"], ["56-1", "56-3"], ["57-1"], ["58-1", "58-2"]]

        def sv(df_r, ests, reg=None):
            if reg and "Regimen" in df_r.columns:
                df_r = df_r[df_r["Regimen"] == reg]
            return [sum(len(df_r[df_r["Estructura"] == e]) for e in g) for g in ests]

        lsd = sv(sin_d, ESTR, "LOSEP"); lm = sv(grp_C, ESTR, "LOSEP"); la = sv(grp_B, ESTR, "LOSEP")
        lsd_t = sum(lsd); lm_t = sum(lm); la_t = sum(la); lt = lsd_t + lm_t + la_t

        ctsd = sv(sin_d, ESTR, "CT"); ctm = sv(grp_C, ESTR, "CT"); cta = sv(grp_B, ESTR, "CT")
        ctsd_t = sum(ctsd); ctm_t = sum(ctm); cta_t = sum(cta); ctt = ctsd_t + ctm_t + cta_t

        esd = sv(sin_d, ESTR, "LOEI"); em = sv(grp_C, ESTR, "LOEI"); ea = sv(grp_B, ESTR, "LOEI")
        esd_t = sum(esd); em_t = sum(em); ea_t = sum(ea); et_ = esd_t + em_t + ea_t

        tsd_t = lsd_t + ctsd_t + esd_t; tm_t = lm_t + ctm_t + em_t; ta_t = la_t + cta_t + ea_t

        res_hdrs = ["DETALLE", "SIN DERECHO", "MENSUALIZA", "ACUMULA", "TOTAL"]
        for ci, h in enumerate(res_hdrs):
            write_hdr(cur_row, ci + 1, h, fill=hdr_fill, font=hdr_font)
        cur_row += 1

        tot_fill = PatternFill("solid", fgColor="D6E4F0")
        for label, sd, m, a, total, fill in [
            ("LOSEP", lsd_t, lm_t, la_t, lt, None),
            ("CÓDIGO DE TRABAJO", ctsd_t, ctm_t, cta_t, ctt, None),
            ("LOEI", esd_t, em_t, ea_t, et_, None),
            ("TOTAL", tsd_t, tm_t, ta_t, gran_total, tot_fill),
        ]:
            write_body(cur_row, 1, label, align=left)
            write_body(cur_row, 2, sd)
            write_body(cur_row, 3, m)
            write_body(cur_row, 4, a)
            write_body(cur_row, 5, total)
            if fill:
                for ci in range(5):
                    ws.cell(row=cur_row, column=ci + 1).fill = fill
                    ws.cell(row=cur_row, column=ci + 1).font = Font(name="Calibri", bold=True, size=8)
            cur_row += 1

        cur_row += 1
        write_hdr(cur_row, 1, "TOTAL FUNCIONARIOS", fill=hdr_fill, font=hdr_font)
        write_hdr(cur_row, 2, str(gran_total), fill=hdr_fill, font=hdr_font)
        cur_row += 2
    else:
        write_hdr(cur_row, 1, "TOTAL FUNCIONARIOS", fill=PatternFill("solid", fgColor="1A3A5C"), font=hdr_font)
        write_hdr(cur_row, 2, str(gran_total), fill=PatternFill("solid", fgColor="1A3A5C"), font=hdr_font)
        cur_row += 2

    wb.save(output_path)
    print(f"✔ Excel de respaldo guardado: {output_path}")


# ═══════════════════════════════════════════════════════════════
# ENSAMBLAJE FINAL — Inyectar en Plantilla.docx
# ═══════════════════════════════════════════════════════════════
def _ensamblar_con_plantilla(doc, plantilla_path, output_path):
    """
    Toma el contenido generado por python-docx (doc),
    extrae su <w:body>, e inyecta en la Plantilla.docx preservando
    headers, footers, logo, styles, etc.
    """
    # 1. Extraer body del doc generado
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    with zipfile.ZipFile(buf, 'r') as z:
        generated_xml = z.read('word/document.xml').decode('utf-8')

    body_match = re.search(r'<w:body>(.*?)<w:sectPr', generated_xml, re.DOTALL)
    body_content = body_match.group(1) if body_match else ''

    # 2. Leer todos los archivos de la plantilla
    with zipfile.ZipFile(plantilla_path, 'r') as z:
        plantilla_files = {item.filename: z.read(item.filename) for item in z.infolist()}

    # 3. Extraer namespace declaration y sectPr de la plantilla
    plantilla_doc_xml = plantilla_files['word/document.xml'].decode('utf-8')
    ns_match   = re.match(r'(<\?xml[^>]+\?>\r?\n?<w:document[^>]+>)', plantilla_doc_xml, re.DOTALL)
    sectPr_match = re.search(r'(<w:sectPr[^>]*>.*?</w:sectPr>)', plantilla_doc_xml, re.DOTALL)
    NS_DECL = ns_match.group(1) if ns_match else ''
    SECT_PR = sectPr_match.group(1) if sectPr_match else ''

    # 4. Construir nuevo document.xml
    new_doc_xml = f'{NS_DECL}<w:body>{body_content}{SECT_PR}</w:body></w:document>'

    # 5. Construir el .docx final copiando todo de la plantilla excepto document.xml
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for fname, data in plantilla_files.items():
            if fname == 'word/document.xml':
                zout.writestr(fname, new_doc_xml.encode('utf-8'))
            else:
                zout.writestr(fname, data)

    with open(output_path, 'wb') as f:
        f.write(out_buf.getvalue())

# ═══════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ═══════════════════════════════════════════════════════════════
def generate(cfg_str, output_path, plantilla_path=None):
    cfg  = json.loads(cfg_str)
    mes  = cfg.get("mes","MES").upper()
    anio = cfg.get("anio","2026")

    # Buscar Plantilla.docx si no se especifica
    if plantilla_path is None:
        plantilla_path = str(Path(__file__).parent / "Plantilla.docx")
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(
            f"No se encontró Plantilla.docx en: {plantilla_path}\n"
            "Asegúrese de que el archivo Plantilla.docx esté en la misma carpeta que app.py"
        )

    print(f"[1/5] Cargando datos...")
    df, sm = load_data(cfg["path_planilla"], cfg["path_matriz"], cfg["path_distributivo"])

    # Guardar datos extra de sin_match (RMU, observación, categoría) para tablas automáticas
    renuncia_extras = {}  # cedula → {rmu, observacion, categoria}
    for m in cfg.get("manual_completados", []):
        ced = norm_ced(m["cedula"]); mask = df["Cedula"]==ced
        # Normalizar estructura: "1-11" → "1--11", "1--11" queda igual
        raw_est = str(m.get("estructura","")).strip()
        if raw_est.upper() in ("1-11","1--11","1 11","1- 11","1-  11"):
            raw_est = "1--11"
        if mask.any():
            df.loc[mask,"Nombre"]     = m["nombre"]
            df.loc[mask,"Estructura"] = raw_est
        renuncia_extras[ced] = {
            "rmu": m.get("rmu", ""),
            "observacion": m.get("observacion", "RENUNCIA / NO CONSTA EN DISTRIBUTIVO"),
            "categoria": m.get("categoria", "RENUNCIA")
        }

    print(f"[2/5] Agrupando {len(df)} registros (clasificación ya aplicada por load_data)...")

    # ── Agrupación basada en MA (calculado en load_data con la nueva tabla de verdad) ──
    # MA=0: Sin Derecho (Derecho=NO, independiente de Solicitud)
    # MA=1: Mensualiza  (Derecho=SI y NO está en Planilla de Pago)
    # MA=2: Acumula     (cédula presente en la Planilla de Pago IESS)

    df["AcumulaFinal"] = df["MA"]

    # Grupos
    gSD = df[df["AcumulaFinal"]==0].copy()   # Sin Derecho
    gB  = df[df["AcumulaFinal"]==2].copy()   # Acumulan
    gC  = df[df["AcumulaFinal"]==1].copy()   # Mensualiza

    # Sub-grupos de Sin Derecho para las nóminas:
    # gA = Sin Derecho con Solicitud SI  (TieneSolicitud=SI / TieneDerecho=NO)
    # gD = Sin Derecho sin Solicitud     (TieneSolicitud=NO / TieneDerecho=NO)
    gA = gSD[gSD["TieneSolicitud"]=="SI"].copy()
    gD = gSD[gSD["TieneSolicitud"]!="SI"].copy()

    print(f"  SD(0):{len(gSD)} [A(SI/NO):{len(gA)} D(NO/NO):{len(gD)}]  Mensualiza(1):{len(gC)}  Acumulan(2):{len(gB)}")
    print(f"  Total: {len(gSD)+len(gC)+len(gB)}")
    sA=split_by_est(gA); sB=split_by_est(gB); sC=split_by_est(gC); sD=split_by_est(gD)

    print(f"[3/5] Construyendo documento Word...")
    doc = Document()
    # Limpiar párrafo vacío que agrega python-docx
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ── DATOS GENERALES ─────────────────────────────────────
    tabla_datos_generales(doc, cfg); spacer(doc)

    # ── 1. ANTECEDENTE ──────────────────────────────────────
    body_para(doc, "1. ANTECEDENTE", bold=True, align="left", before=100)
    body_para(doc, "El Código del Trabajo (CT) y la Ley del Servidor Público (LOSEP) establecen que el empleador debe pagar mensualmente a sus empleados, por fondo de reserva, el 8,33% del salario o remuneración que percibe el empleado o trabajador, a partir del segundo año de trabajo en la misma entidad.")
    body_para(doc, "Desde de 2009, con las reformas a la Ley de Seguridad Social, el afiliado al IESS es quien decide si el fondo de reserva lo recibe juntamente con su salario mensual o si lo ahorra en el IESS.")
    body_para(doc, "El pago de fondo de reserva es un derecho irrenunciable de todos los trabajadores en relación de dependencia, que se hace efectivo después de que ha cumplido el primer año de trabajo para un mismo empleador, según lo establece el capítulo 11, artículo 196, del Código del Trabajo; y el Art. 269 del Reglamento de la LOSEP.")
    body_para(doc, "El plazo que tiene el empleador para depositar el fondo de reserva en el IESS es el mismo que tiene para pagar los aportes: entre los 15 primeros días del mes siguiente al trabajado. Si no cancela ese monto dentro de este plazo, el empleador cae en mora y tiene que pagar multas más intereses, de acuerdo con la Ley de Seguridad Social.")

    # ── 2. BASE LEGAL ────────────────────────────────────────
    body_para(doc, "2. BASE LEGAL", bold=True, align="left", before=100)
    body_para(doc, "1.    LEY ORGANICA DEL SERVICIO PUBLICO", bold=True, align="left")
    body_para(doc, "\u201cArt. 99.- Fondos de Reserva. - Los servidores y servidoras de las entidades, instituciones, organismos o personas jurídicas señaladas en el artículo 3 de esta ley, tienen derecho a recibir anualmente y a partir del segundo año por concepto de fondos de reserva una remuneración mensual unificada del servidor equivalente a la que perciba, conforme a las normas pertinentes que regulan la seguridad social.\u201d", italic=True)
    body_para(doc, "2.    REGLAMENTO A LA LEY ORGANICA DEL SERVICIO PUBLICO", bold=True, align="left")
    body_para(doc, "Art. 269.- Del fondo de reserva. - En caso de que una servidora o un servidor cesare en funciones en una de las instituciones establecidas en el artículo 3 de la LOSEP, e ingrese al primer día laborable siguiente a otra institución, entidad u organismo del sector público, no perderá su derecho y antigüedad para el cálculo, provisión y pago del fondo de reserva.", italic=True)
    body_para(doc, "2.3 REGLAMENTO PARA EL PAGO O DEVOLUCION DE LOS FONDOS DE RESERVA", bold=True, align="left")
    body_para(doc, "\u201cArt. 1.- Derecho del trabajador o servidor. - El trabajador o servidor público con relación de dependencia, tendrá derecho al pago mensual del Fondo de Reserva por parte de su empleador, en un porcentaje equivalente al ocho coma treinta y tres por ciento (8,33%) de la remuneración aportada al Instituto Ecuatoriano de Seguridad Social, después del primer año de trabajo.\u201d", italic=True)
    body_para(doc, "\u201cArt. 2.- Depósito del Fondo de Reserva en el IESS. - Cuando el trabajador o servidor público solicitare por escrito, a través del aplicativo informático, que su Fondo de Reserva sea depositado en el IESS, el empleador consignará dicho rubro mensual, juntamente con la planilla de aportes. Estos valores se registrarán en la cuenta individual del afiliado para su posterior devolución. De no depositarse el Fondo de Reserva dentro de los primeros quince (15) días del mes siguiente al que corresponda, causará mora con los recargos y multas correspondientes.\u201d", italic=True)
    body_para(doc, "3.    CODIGO DE TRABAJO", bold=True, align="left")
    body_para(doc, "\u201cArt. 196.- Derecho al fondo de reserva. Todo trabajador que preste servicios por más de un año tiene derecho a que el empleador le abone una suma equivalente a un mes de sueldo o salario por cada año completo posterior al primero de sus servicios. Estas sumas constituirán su fondo de reserva o trabajo capitalizado. El trabajador no perderá este derecho por ningún motivo. La determinación de la cantidad que corresponda por cada año de servicio se hará de acuerdo con lo dispuesto en el artículo 95 de este Código.\u201d", italic=True)

    # ── 3. OBJETIVO ──────────────────────────────────────────
    body_para(doc, "3. OBJETIVO", bold=True, align="left", before=100)
    body_para(doc, f"Determinar los valores a cancelar y depositar en el IESS, por concepto de Fondos de Reserva al personal que labora en la Dirección Distrital de Educación Los Chillos, correspondiente al mes de {mes} DE {anio}.")

    # ── 4. DESARROLLO ────────────────────────────────────────
    body_para(doc, "4. DESARROLLO", bold=True, align="left", before=100)
    body_para(doc, f"Una vez que la Unidad Distrital de Talento Humano ha realizado la revisión y el cruce de valores de la planilla de fondos de reserva correspondiente al MES DE {mes} DE {anio} misma que fue remitida mediante correo electrónico de fecha {cfg.get('fecha_remision','')}, por Cecilia Acurio, Analista Distrital Administrativo Financiero, con el distributivo de trabajo, con la información recibida se prepara el informe de FONDOS DE RESERVA del MES DE {mes} DE {anio}, en el que se desprende el siguiente detalle:")

    # ── NOVEDADES DINÁMICAS (numeradas a partir del ítem 5) ──
    print(f"[4/5] Insertando novedades y nóminas...")
    novedades = cfg.get("novedades", [])

    # Contador de ítems: 1=Antecedente, 2=Base Legal, 3=Objetivo, 4=Desarrollo
    # Las novedades arrancan en 5
    item_num = 5

    for nov in novedades:
        titulo_base = nov.get("titulo", "NOVEDAD").strip()
        headers = nov.get("headers", [])
        widths  = nov.get("widths", [])
        filas   = nov.get("filas", [])
        # Asignar número de ítem automáticamente
        titulo_con_num = f"{item_num}.- NOVEDAD: {titulo_base}" if not titulo_base.startswith(str(item_num)) else titulo_base
        # Evitar doble "NOVEDAD: NOVEDAD:"
        if "NOVEDAD:" in titulo_con_num and titulo_base.upper().startswith("NOVEDAD:"):
            titulo_con_num = f"{item_num}.- {titulo_base}"
        heading_para(doc, titulo_con_num, underline=True)
        if headers and widths:
            tabla_novedad(doc, headers, widths, filas)
        spacer(doc)
        item_num += 1

    # ── ÍTEMS AUTOMÁTICOS POR CATEGORÍA (sin_match agrupados) ──
    # sm contiene los registros de la Matriz que NO están en el Distributivo
    sm_ceds = {norm_ced(m["Cedula"]) for m in sm}
    renuncias_df = df[df["Cedula"].isin(sm_ceds)]

    # Agrupar por categoría asignada en el paso 4
    CATEGORIAS_SM = ["RENUNCIA", "NOTIFICACIÓN DE SALIDA", "FALLECIMIENTO"]
    for cat in CATEGORIAS_SM:
        cat_rows = []
        for _, row in renuncias_df.iterrows():
            ced = row["Cedula"]
            extras = renuncia_extras.get(ced, {})
            row_cat = extras.get("categoria", "RENUNCIA")
            if row_cat == cat:
                cat_rows.append([
                    "",  # auto-numerado por tabla_novedad
                    ced,
                    row["Nombre"] if row["Nombre"] and str(row["Nombre"]) != "nan" else "",
                    extras.get("rmu", ""),
                    extras.get("observacion", f"{cat} / NO CONSTA EN DISTRIBUTIVO")
                ])
        if cat_rows:
            heading_para(doc, f"{item_num}.- NOVEDAD: {cat}", underline=True)
            hdrs_r = ["No", "CÉDULA", "APELLIDOS Y NOMBRES", "RMU", "OBSERVACIÓN"]
            wids_r = [300, 900, 2500, 700, 3247]
            tabla_novedad(doc, hdrs_r, wids_r, cat_rows)
            spacer(doc)
            item_num += 1

    # El número de ítem siguiente es para la nómina
    num_nomina = item_num

    # ── NÓMINAS ──────────────────────────────────────────────
    body_para(doc, f"{num_nomina}. NOMINA DE FUNCIONARIOS QUE NO TIENEN DERECHO A COBRAR FONDOS DE RESERVA",
              bold=True, align="left", before=100)

    heading_para(doc, f"{num_nomina}.1.- TIENE SOLICITUD SI / TIENE DERECHO NO", underline=True)
    for est, sub in sA.items(): tabla_nomina(doc, sub, 0); spacer(doc)

    heading_para(doc, f"{num_nomina}.1.- TIENE SOLICITUD NO / TIENE DERECHO NO", underline=True)
    for est, sub in sD.items(): tabla_nomina(doc, sub, 0); spacer(doc)

    body_para(doc, f"{num_nomina}.2. NOMINA DE FUNCIONARIOS QUE TIENEN SOLICITUD SI / TIENEN DERECHO",
              bold=True, align="left")
    heading_para(doc, "ACUMULAN", underline=True)
    for est, sub in sB.items(): tabla_nomina(doc, sub, 2); spacer(doc)

    body_para(doc, f"{num_nomina}.3. TIENEN SOLICITUD NO / TIENEN DERECHO SI", bold=True, align="left")
    heading_para(doc, "MENSUALIZA", underline=True)
    for est, sub in sC.items(): tabla_nomina(doc, sub, 1); spacer(doc)

    # ── CONCLUSIÓN ───────────────────────────────────────────
    num_conclu = num_nomina + 1
    body_para(doc, f"{num_conclu}. CONCLUSIÓN", bold=True, align="left", before=100)
    body_para(doc, f"En base a lo anteriormente citado y de conformidad a leyes y reglamentos vigentes para el pago de Fondos de Reserva, la Unidad Distrital de Talento Humano solicita a la División Distrital Financiera, realizar la cancelación de la planilla de fondos de reserva correspondiente al mes de {mes} {anio} según el siguiente detalle:")

    gran = tabla_resumen(doc, gA, gB, gC, gD); spacer(doc)

    body_para(doc, "RECOMENDACIÓN", bold=True, align="left", before=100)
    body_para(doc, f"La Unidad Distrital de Talento Humano de la DIRECCIÓN DISTRITAL 17D08-PARROQUIAS RURALES: CONOCOTO A LA MERCED \u2013EDUCACIÓN, acogiéndose a lo determinado en las leyes y reglamentos citados anteriormente, recomienda que el Director Distrital, como autoridad nominadora luego de conocer el presente informe autorice a la Unidad Administrativa Financiera, realizar el trámite correspondiente para cancelar los valores de fondos de reserva del mes {mes} {anio} con {gran} funcionarios que laboran en la Dirección Distrital 17D08 Parroquias Rurales Conocoto a la Merced - Educación, de conformidad con el detalle descrito en el cuadro anterior.")
    spacer(doc)

    tabla_firmas(doc, cfg)

    print(f"[5/5] Ensamblando con plantilla y guardando...")
    _ensamblar_con_plantilla(doc, plantilla_path, output_path)

    # Generar Excel de respaldo con todas las tablas
    excel_path = str(output_path).replace('.docx', '_TABLAS.xlsx')
    try:
        generar_excel_respaldo(cfg, novedades, sA, sB, sC, sD, gran, excel_path,
                              grp_A=gA, grp_B=gB, grp_C=gC, grp_D=gD,
                              renuncias_df=renuncias_df, renuncia_extras=renuncia_extras)
    except Exception as e:
        print(f"  ⚠ No se pudo generar Excel de respaldo: {e}")
        excel_path = None

    print(f"\u2714 Listo \u2014 {gran} funcionarios \u2014 {output_path}")
    return gran, sm, excel_path


if __name__ == "__main__":
    f = sys.argv[1] if len(sys.argv) > 1 else "config.json"
    with open(f) as fh: cfg_str = fh.read()
    cfg = json.loads(cfg_str)
    out = cfg.get("output_path", "/tmp/informe.docx")
    total, sm, excel_path = generate(cfg_str, out)
    print(json.dumps({"ok": True, "total": total, "sin_match": sm, "output": out, "excel": excel_path}))
